#!/usr/bin/env python3
"""Build a Word document, applying the Clear Sky corporate template styling.

When a template is provided, the document is generated *from* that template so
its named styles (CS Eyebrow, CS Page Title, CS Deck, Heading 1, Body Text,
CS Table Header/Body, ...), theme colors, and fonts are used directly.
"""

from __future__ import annotations

import argparse
import json
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# Clear Sky table palette (from the style guide).
HEADER_FILL = "061A2E"    # Midnight Navy — table header fill
HEADER_TEXT = "FFFFFF"    # White — table header text
ROW_MIST = "F3F6F8"       # Mist — alternating body-row fill
ROW_WHITE = "FFFFFF"      # White — alternating body-row fill
ROW_RULE = "00758F"       # System Teal — rules / small accents
KEYWORD = "00758F"        # System Teal (accent2) — Gherkin keyword emphasis

STEP_FONT_SIZE = 14       # Step text — larger for readability against screenshots

# Decision-table geometry from the template (twips / dxa).
TABLE_WIDTH = 9360
CELL_MARGIN_TB = 100
CELL_MARGIN_LR = 120
BODY_CELL_MARGIN_TB = 95

GHERKIN_KEYWORDS = {"Given", "When", "Then", "And", "But"}


def has_style(doc, name: str) -> bool:
    try:
        return name in (s.name for s in doc.styles)
    except Exception:
        return False


def add_run(paragraph, text: str, *, bold=False, italic=False, size=None, mono=False, color=None):
    run = paragraph.add_run(text)
    if mono:
        run.font.name = "Courier New"
        rpr = run._element.get_or_add_rPr()
        rpr.rFonts.set(qn("w:eastAsia"), "Courier New")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic
    return run


def styled_paragraph(doc, style: str, fallback: str = "Normal"):
    """Add a paragraph using `style` when the template defines it."""
    return doc.add_paragraph(style=style if has_style(doc, style) else fallback)


def keep_cover_only(doc):
    """Remove template placeholder pages but keep the cover page + section break.

    Returns the final (content) sectPr element so the caller can configure it.
    The body ends up as: [cover table, cover section-break paragraph, <content>, final sectPr].
    """
    body = doc.element.body
    children = list(body)

    final_sectpr = children[-1] if children and children[-1].tag == qn("w:sectPr") else None
    cover_table = None
    cover_sect_para = None
    for ch in children:
        if ch.tag == qn("w:tbl") and cover_table is None:
            cover_table = ch
        elif ch.tag == qn("w:p"):
            p_pr = ch.find(qn("w:pPr"))
            if p_pr is not None and p_pr.find(qn("w:sectPr")) is not None and cover_sect_para is None:
                cover_sect_para = ch

    keep = {id(el) for el in (cover_table, cover_sect_para, final_sectpr) if el is not None}
    for ch in children:
        if id(ch) not in keep:
            body.remove(ch)

    return cover_table, final_sectpr


def configure_content_section(doc, final_sectpr) -> None:
    """Content section shows the branded header/footer on every page (page 2 onward).

    Removing titlePg makes Word use the default (branded) header/footer on the
    section's first page too, so header + footer effectively start after the cover.
    """
    if final_sectpr is not None:
        for title_pg in final_sectpr.findall(qn("w:titlePg")):
            final_sectpr.remove(title_pg)
    # Set readable content margins (cover section keeps its full-bleed 0 margins).
    section = doc.sections[-1]
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)


def replace_placeholders(element, mapping: dict) -> None:
    """Replace placeholder strings in a story, tolerating text split across runs.

    Paragraphs containing images or fields (e.g. the logo or PAGE number) are only
    edited run-by-run so their drawings/fields are never disturbed.
    """
    for para in element.iter(qn("w:p")):
        has_drawing = para.find(".//" + qn("w:drawing")) is not None
        has_field = (
            para.find(".//" + qn("w:fldChar")) is not None
            or para.find(".//" + qn("w:instrText")) is not None
        )
        text_nodes = [t for run in para.findall(qn("w:r")) for t in run.findall(qn("w:t"))]
        if not text_nodes:
            continue

        if has_drawing or has_field:
            for t in text_nodes:
                if t.text:
                    for key, value in mapping.items():
                        if key in t.text:
                            t.text = t.text.replace(key, value)
            continue

        joined = "".join(t.text or "" for t in text_nodes)
        replaced = joined
        for key, value in mapping.items():
            if key in replaced:
                replaced = replaced.replace(key, value)
        if replaced != joined:
            text_nodes[0].text = replaced
            for t in text_nodes[1:]:
                t.text = ""


def _set_cell_text(cell_tc, text: str) -> None:
    para = cell_tc.find(qn("w:p"))
    if para is None:
        return
    t_nodes = [t for run in para.findall(qn("w:r")) for t in run.findall(qn("w:t"))]
    if not t_nodes:
        return
    t_nodes[0].text = text
    for t in t_nodes[1:]:
        t.text = ""


def _relabel_or_remove(cover_table, label: str, new_label: str, value: str) -> None:
    """The cover metadata is a nested [label | value] table.

    Reuse the row for `new_label`/`value` when a value is present; otherwise drop the row.
    """
    for row in cover_table.iter(qn("w:tr")):
        cells = row.findall(qn("w:tc"))
        if not cells:
            continue
        first_text = "".join(t.text or "" for t in cells[0].iter(qn("w:t"))).strip()
        if first_text != label:
            continue
        if value:
            _set_cell_text(cells[0], new_label)
            if len(cells) > 1:
                _set_cell_text(cells[1], value)
        else:
            row.getparent().remove(row)
        return


def fill_cover(cover_table, title: str, description: str, tags: list, feature: str, test_suite: str) -> None:
    if cover_table is None:
        return
    subtitle = description or title
    tags_line = ("Tags: " + ", ".join(tags)) if tags else ""
    mapping = {
        "CORPORATE COMMUNICATIONS TEMPLATE": "SCENARIO DOCUMENTATION",
        "[DOCUMENT TITLE]": title,
        "[Document subtitle or outcome-oriented descriptor]": subtitle,
        "Professional reports, proposals, strategy documents, executive updates, "
        "and external communications": tags_line,
        "[Author name]": "Clear Sky Solutions",
        "[Month DD, YYYY]": date.today().strftime("%B %d, %Y"),
        "[CONFIDENTIAL / INTERNAL / PUBLIC]": "INTERNAL",
    }
    replace_placeholders(cover_table, mapping)

    # Reuse the DEPARTMENT / VERSION metadata slots for Feature / Test Suite when
    # present in the JSON; otherwise remove them (per "take out department, version").
    _relabel_or_remove(cover_table, "DEPARTMENT", "FEATURE", feature)
    _relabel_or_remove(cover_table, "VERSION", "TEST SUITE", test_suite)


def fill_headers_footers(doc, title: str) -> None:
    mapping = {
        "[DOCUMENT TITLE]": title,
        # No status in the header — it already appears in the footer (bottom-left).
        "[STATUS]": "",
        "[CONFIDENTIALITY / DOCUMENT STATUS]": "INTERNAL",
    }
    for rel in doc.part.rels.values():
        if "header" in rel.reltype or "footer" in rel.reltype:
            replace_placeholders(rel.target_part.element, mapping)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    for existing in tc_pr.findall(qn("w:shd")):
        tc_pr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_bottom_border(cell, color: str = ROW_RULE, size: str = "4") -> None:
    """Body cells use a single thin teal bottom rule (per the 'thin teal rules' rule)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    for existing in tc_pr.findall(qn("w:tcBorders")):
        tc_pr.remove(existing)
    borders = OxmlElement("w:tcBorders")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "0")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    tc_pr.append(borders)


def set_cell_margins(cell, top: int, bottom: int, left: int | None = None, right: int | None = None) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    for existing in tc_pr.findall(qn("w:tcMar")):
        tc_pr.remove(existing)
    mar = OxmlElement("w:tcMar")
    for edge, value in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        if value is None:
            continue
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:w"), str(value))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tc_pr.append(mar)


def set_cell_width(cell, width: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    for existing in tc_pr.findall(qn("w:tcW")):
        tc_pr.remove(existing)
    tc_w = OxmlElement("w:tcW")
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")
    tc_pr.insert(0, tc_w)


def clear_table_borders(table) -> None:
    """Remove the default grid so only the template's bottom rules remain."""
    tbl_pr = table._tbl.tblPr
    for existing in tbl_pr.findall(qn("w:tblBorders")):
        tbl_pr.remove(existing)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    tbl_pr.append(borders)


def apply_decision_table_layout(table) -> None:
    """Apply the Clear Sky decision-table geometry from the corporate template."""
    tbl_pr = table._tbl.tblPr

    for tag in ("tblW", "jc", "tblLayout", "tblCellMar", "tblLook"):
        for existing in tbl_pr.findall(qn(f"w:{tag}")):
            tbl_pr.remove(existing)

    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_w)

    # Rules say "left-aligned grids" (do not replicate the template's centered example).
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "left")
    tbl_pr.append(jc)

    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)

    cell_mar = OxmlElement("w:tblCellMar")
    for edge, value in (
        ("top", CELL_MARGIN_TB),
        ("left", CELL_MARGIN_LR),
        ("bottom", CELL_MARGIN_TB),
        ("right", CELL_MARGIN_LR),
    ):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:w"), str(value))
        el.set(qn("w:type"), "dxa")
        cell_mar.append(el)
    tbl_pr.append(cell_mar)

    look = OxmlElement("w:tblLook")
    look.set(qn("w:val"), "04A0")
    look.set(qn("w:firstRow"), "1")
    look.set(qn("w:lastRow"), "0")
    look.set(qn("w:firstColumn"), "1")
    look.set(qn("w:lastColumn"), "0")
    look.set(qn("w:noHBand"), "0")
    look.set(qn("w:noVBand"), "1")
    tbl_pr.append(look)

    clear_table_borders(table)


def add_step_title(doc, text: str) -> None:
    """Body Text paragraph; the leading Gherkin keyword is uppercase, bold, deep blue.

    GIVEN/WHEN/THEN sit flush left; AND/BUT are indented.
    """
    p = styled_paragraph(doc, "Body Text")
    p.paragraph_format.space_before = Pt(6)
    parts = text.split(" ", 1)
    if len(parts) == 2 and parts[0] in GHERKIN_KEYWORDS:
        keyword = parts[0].upper()
        if keyword in ("AND", "BUT"):
            p.paragraph_format.left_indent = Inches(0.35)
        add_run(p, keyword, bold=True, color=KEYWORD, size=STEP_FONT_SIZE)
        add_run(p, " " + parts[1], size=STEP_FONT_SIZE)
    else:
        add_run(p, text, size=STEP_FONT_SIZE)


def add_data_table(doc, table_rows) -> None:
    """Build a data table matching the template's decision-table layout."""
    cols = len(table_rows[0])
    table = doc.add_table(rows=len(table_rows), cols=cols)
    table.style = "Normal Table"
    apply_decision_table_layout(table)

    col_width = TABLE_WIDTH // cols
    header_style = "CS Table Header" if has_style(doc, "CS Table Header") else None
    body_style = "CS Table Body" if has_style(doc, "CS Table Body") else None

    for r_idx, row in enumerate(table_rows):
        is_header = r_idx == 0
        # Alternating body rows: white, mist, white, mist...
        body_fill = ROW_WHITE if (r_idx % 2 == 1) else ROW_MIST

        for c_idx, cell_value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            set_cell_width(cell, col_width)

            para = cell.paragraphs[0]
            if is_header:
                if header_style:
                    para.style = header_style
                else:
                    add_run(para, str(cell_value), bold=True, color=HEADER_TEXT, size=9)
                    set_cell_shading(cell, HEADER_FILL)
                    continue
                para.add_run(str(cell_value))
                set_cell_shading(cell, HEADER_FILL)
            else:
                if body_style:
                    para.style = body_style
                    para.add_run(str(cell_value))
                else:
                    add_run(para, str(cell_value), size=9)
                set_cell_shading(cell, body_fill)
                set_cell_bottom_border(cell)
                set_cell_margins(cell, BODY_CELL_MARGIN_TB, BODY_CELL_MARGIN_TB)


def build_doc(manifest: dict, output: Path, template: Path | None) -> None:
    title = (manifest.get("title") or "Generated Scenario").strip()
    description = (manifest.get("description") or "").strip()
    tags = [str(tag).strip() for tag in (manifest.get("tags") or []) if str(tag).strip()]
    feature = (manifest.get("feature") or "").strip()
    test_suite = (manifest.get("testSuite") or "").strip()

    has_cover = False
    if template and template.exists():
        doc = Document(str(template))
        cover_table, final_sectpr = keep_cover_only(doc)
        fill_cover(cover_table, title, description, tags, feature, test_suite)
        configure_content_section(doc, final_sectpr)
        fill_headers_footers(doc, title)
        has_cover = cover_table is not None
    else:
        doc = Document()

    # Without the template cover, print the title block on the first content page.
    if not has_cover:
        styled_paragraph(doc, "CS Eyebrow").add_run("SCENARIO")
        styled_paragraph(doc, "CS Page Title", fallback="Title").add_run(title)
        if description:
            styled_paragraph(doc, "CS Deck", fallback="Subtitle").add_run(description)
        if tags:
            p = styled_paragraph(doc, "CS Small")
            add_run(p, "Tags: ", bold=True)
            add_run(p, ", ".join(tags))
        if feature:
            p = styled_paragraph(doc, "CS Small")
            add_run(p, "Feature: ", bold=True)
            add_run(p, feature)
        if test_suite:
            p = styled_paragraph(doc, "CS Small")
            add_run(p, "Test Suite: ", bold=True)
            add_run(p, test_suite)

    doc.add_paragraph(style="Heading 1" if has_style(doc, "Heading 1") else "Normal").add_run("Steps")

    steps = manifest.get("steps") or []
    if not steps:
        styled_paragraph(doc, "Body Text").add_run("No steps were found in this export.")

    for index, step in enumerate(steps):
        step_text = step.get("text") or "(Untitled step)"
        add_step_title(doc, step_text)

        table_rows = step.get("table") or []
        screenshot_paths = step.get("screenshotPaths") or []
        if table_rows:
            add_data_table(doc, table_rows)
            if screenshot_paths:
                doc.add_paragraph("")

        for shot_path in screenshot_paths:
            path = Path(shot_path)
            if not path.exists():
                continue
            picture = doc.add_paragraph()
            picture.alignment = WD_ALIGN_PARAGRAPH.LEFT
            picture.add_run().add_picture(str(path), width=Inches(6.0))

        # Blank line between steps
        if index < len(steps) - 1:
            doc.add_paragraph("")

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("manifest")
    parser.add_argument("--out", required=True)
    parser.add_argument("--template", default=None)
    args = parser.parse_args()
    manifest = json.loads(Path(args.manifest).read_text())
    template = Path(args.template) if args.template else None
    build_doc(manifest, Path(args.out), template)
    print(f"Generated {args.out}")


if __name__ == "__main__":
    main()
